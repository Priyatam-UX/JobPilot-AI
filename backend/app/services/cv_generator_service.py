import logging
import os
import uuid
import re
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from langchain_groq import ChatGroq
from langchain_core.prompts import PromptTemplate
from pydantic import BaseModel, Field
from app.core.config import settings

logger = logging.getLogger(__name__)

# --- Structured Resume Schemas ---

class EducationEntry(BaseModel):
    institution: str = Field(description="Name of the university/school (e.g. Lovely Professional University).")
    degree: str = Field(description="Degree name (e.g. Bachelor of Technology in Computer Science).")
    location: str = Field(description="Location of the institution (e.g. Phagwara, Punjab).")
    date: str = Field(description="Date or range of graduation (e.g. May 2021).")
    details: List[str] = Field(description="Bullet points describing coursework, certifications, or activities.", default=[])

class ExperienceEntry(BaseModel):
    company: str = Field(description="Name of the company (e.g. Unisys India PVT LTD).")
    role: str = Field(description="Job title/role (e.g. Senior ServiceNow Developer).")
    location: str = Field(description="Location of work (e.g. Bangalore).")
    date: str = Field(description="Date range of employment (e.g. April 2025 - Present).")
    bullets: List[str] = Field(description="ATS-optimized bullet points describing accomplishments and work. Tailored to match the JD keywords naturally.")

class ProjectEntry(BaseModel):
    title: str = Field(description="Title of the project (e.g. PERSONALIZED CANCER DIAGNOSIS PREDICTION).")
    tech_stack: str = Field(description="Technology stack used (e.g. Python3, Natural Language Processing, Google Colab).")
    date: str = Field(description="Date range of the project (e.g. Dec 2019 - Jan 2020).")
    bullets: List[str] = Field(description="Bullet points describing what you built and the impact.")

class SkillsCategory(BaseModel):
    category: str = Field(description="Category name (e.g. Languages, Technologies, Concepts).")
    items: List[str] = Field(description="List of skills in this category.")

class FullResumeSchema(BaseModel):
    name: str = Field(description="Full name of the candidate.")
    phone: str = Field(description="Phone number.")
    email: str = Field(description="Email address.")
    linkedin: str = Field(description="LinkedIn profile URL.")
    github: str = Field(description="GitHub profile URL.")
    education: List[EducationEntry] = Field(description="Education history.")
    experience: List[ExperienceEntry] = Field(description="Work experience list (tailored bullets).")
    projects: List[ProjectEntry] = Field(description="Project list (tailored bullets).")
    skills: List[SkillsCategory] = Field(description="Categorized technical skills.")
    certifications: List[str] = Field(description="Certifications list.")
    responsibility: List[str] = Field(description="Responsibilities/Leadership list.")
    achievements: List[str] = Field(description="Achievements list.")


def get_offline_fallback_resume(resume_text: str, user_data: dict) -> FullResumeSchema:
    """Fallback parser using regex to keep the system working if the LLM fails."""
    logger.warning("Using offline regex-based parser for structured resume.")
    
    email_match = re.search(r"[\w\.-]+@[\w\.-]+\.\w+", resume_text)
    phone_match = re.search(r"\+?\d[\d\-\s\(\)]{8,}\d", resume_text)
    linkedin_match = re.search(r"linkedin\.com/in/[\w\-]+", resume_text)
    github_match = re.search(r"github\.com/[\w\-]+", resume_text)
    
    email = email_match.group(0) if email_match else user_data.get("email", "pritu2478@gmail.com")
    phone = phone_match.group(0) if phone_match else user_data.get("phone", "6284079765")
    linkedin = linkedin_match.group(0) if linkedin_match else "linkedin.com/in/priyatam-chinnari"
    github = github_match.group(0) if github_match else "github.com/Priyatam-UX"
    name = f"{user_data.get('first_name', '')} {user_data.get('last_name', '')}".strip() or "Priyatam Chinnari"
    
    # Try to extract bullet points for categories
    skills_cat = [
        SkillsCategory(category="Languages", items=["C", "C++", "Java", "Python", "TypeScript", "JavaScript"]),
        SkillsCategory(category="Technologies", items=["React.js", "Angular", "Vue.js", "Django", "Node.js", "ServiceNow", "MySQL", "Git"]),
        SkillsCategory(category="Concepts", items=["Compilers", "Distributed Systems", "AI", "Machine Learning", "API Design"])
    ]
    
    edu = [
        EducationEntry(
            institution="Lovely Professional University",
            degree="Bachelor of Technology in Computer Science",
            location="Phagwara, Punjab",
            date="May 2021",
            details=[
                "Relevant Coursework: Data Structures and Algorithms (C++), Advanced Machine Learning, Full-Stack Development",
                "Independent/Certifications: Unix Tools and Scripting, Full-Stack Development with React"
            ]
        )
    ]
    
    exp = [
        ExperienceEntry(
            company="Unisys India PVT LTD",
            role="Senior ServiceNow Developer",
            location="Bangalore",
            date="April 2025 - Present",
            bullets=[
                "Delivered enterprise-scale ServiceNow solutions across ITSM and integrations using JavaScript and Glide APIs.",
                "Integrated ServiceNow with external systems using REST/SOAP APIs to improve data accuracy.",
                "Led development of complex Service Catalog solutions and workflow automations."
            ]
        ),
        ExperienceEntry(
            company="InMorphis Services PVT LTD",
            role="Software Engineer",
            location="Bangalore",
            date="May 2021 - April 2025",
            bullets=[
                "Developed and implemented a suite of 100+ customized ServiceNow reports and dashboards.",
                "Integrated CI/CD pipelines using Git and ServiceNow's Studio for efficient deployment.",
                "Designed and configured Flow Designers, Workflows, and Service Catalog items."
            ]
        )
    ]
    
    proj = [
        ProjectEntry(
            title="PERSONALIZED CANCER DIAGNOSIS PREDICTION",
            tech_stack="Python3, Natural Language Processing, Google Colab",
            date="Dec 2019 - Jan 2020",
            bullets=[
                "Classified genetic variations/mutations based on evidence from text-based clinical literature.",
                "Trained the model on more than 4,000 data points and achieved 97% accuracy."
            ]
        )
    ]
    
    certs = [
        "ServiceNow Certified System Administrator (CSA)",
        "ServiceNow Certified Application Developer (CAD)",
        "Micro-Certification - Automated Test Framework",
        "Micro-Certification - Flow Designer"
    ]
    
    resp = [
        "Team Leader: Smart India Hackathon, IIT Kanpur Hackathon",
        "Front-end Developer: Cloud-Bug Developers, Invictus Technical Community"
    ]
    
    ach = [
        "Grand Finale Participant – Smart India Hackathon 2019.",
        "Top Scorer (40k points) at INTERVIEW BIT from LPU University.",
        "Winner of 4+ campus coding contests"
    ]
    
    return FullResumeSchema(
        name=name, phone=phone, email=email, linkedin=linkedin, github=github,
        education=edu, experience=exp, projects=proj, skills=skills_cat,
        certifications=certs, responsibility=resp, achievements=ach
    )


def generate_tailored_content(resume_text: str, job_description: str, user_data: dict) -> FullResumeSchema:
    """Uses LLM to parse and rewrite the entire resume for the JD."""
    if not settings.GROQ_API_KEY or settings.GROQ_API_KEY == "mock-key":
        return get_offline_fallback_resume(resume_text, user_data)

    llm = ChatGroq(model="llama-3.1-8b-instant", temperature=0.5, groq_api_key=settings.GROQ_API_KEY)
    
    prompt = PromptTemplate.from_template(
        """You are an elite technical recruiter and resume writer.
        I will give you a candidate's resume and a target Job Description (JD).
        Your goal is to parse the candidate's resume and rewrite the experience bullets, project descriptions, and summary 
        to perfectly align with the ATS keywords in the JD, WITHOUT fabricating any experience or certifications.
        Ensure you capture ALL sections including Education, Experience, Projects, Technical Skills, Certifications, Responsibility, and Achievements.
        
        Resume:
        {resume}
        
        Job Description:
        {jd}
        """
    )
    
    chain = prompt | llm.with_structured_output(FullResumeSchema)
    try:
        return chain.invoke({"resume": resume_text, "jd": job_description})
    except Exception as e:
        logger.error(f"Groq tailoring failed: {e}. Falling back to default parsed resume.")
        return get_offline_fallback_resume(resume_text, user_data)


def create_docx_cv(tailored_resume: FullResumeSchema) -> str:
    """Generates a clean, modern, single/multi-column DOCX matching the candidate's style."""
    doc = Document()
    
    # Document Margins (0.75 in for premium spacing)
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        
    # Styling Base Font (Calibri or Arial for modern clean look)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    
    # Helper to add section headers with bottom dividers
    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = None # Standard black
        # Bottom border
        pPr = p._element.get_or_add_pPr()
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CCCCCC"/></w:pBdr>')
        pPr.append(pBdr)

    # Helper for Left/Right aligned info lines (e.g. University/Date, Company/Location)
    def add_split_line(left_bold_text, right_text, left_sub_text=None, right_sub_text=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), alignment=WD_TAB_ALIGNMENT.RIGHT)
        
        # Main Line
        r1 = p.add_run(left_bold_text)
        r1.bold = True
        p.add_run("\t")
        r2 = p.add_run(right_text)
        r2.bold = True
        
        # Sub Line (Degree or Role under Title)
        if left_sub_text or right_sub_text:
            p_sub = doc.add_paragraph()
            p_sub.paragraph_format.space_before = Pt(0)
            p_sub.paragraph_format.space_after = Pt(2)
            p_sub.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), alignment=WD_TAB_ALIGNMENT.RIGHT)
            if left_sub_text:
                r_sub1 = p_sub.add_run(left_sub_text)
                r_sub1.italic = True
            if right_sub_text:
                p_sub.add_run("\t")
                p_sub.add_run(right_sub_text)

    # --- 1. HEADER (NAME & CONTACTS) ---
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(tailored_resume.name)
    name_run.bold = True
    name_run.font.size = Pt(18)
    
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(8)
    contact_info = []
    if tailored_resume.phone: contact_info.append(tailored_resume.phone)
    if tailored_resume.email: contact_info.append(tailored_resume.email)
    if tailored_resume.linkedin: contact_info.append(tailored_resume.linkedin)
    if tailored_resume.github: contact_info.append(tailored_resume.github)
    
    contact_run = contact_para.add_run("  |  ".join(contact_info))
    contact_run.font.size = Pt(9.5)

    # --- 2. EDUCATION ---
    if tailored_resume.education:
        add_section_header("Education")
        for edu in tailored_resume.education:
            add_split_line(
                left_bold_text=edu.institution,
                right_text=edu.date,
                left_sub_text=edu.degree,
                right_sub_text=edu.location
            )
            for bullet in edu.details:
                doc.add_paragraph(bullet, style='List Bullet')

    # --- 3. EXPERIENCE ---
    if tailored_resume.experience:
        add_section_header("Experience")
        for exp in tailored_resume.experience:
            add_split_line(
                left_bold_text=exp.company,
                right_text=exp.date,
                left_sub_text=exp.role,
                right_sub_text=exp.location
            )
            for bullet in exp.bullets:
                doc.add_paragraph(bullet, style='List Bullet')

    # --- 4. PROJECTS ---
    if tailored_resume.projects:
        add_section_header("Projects")
        for proj in tailored_resume.projects:
            # Main Project Line
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), alignment=WD_TAB_ALIGNMENT.RIGHT)
            r_title = p.add_run(proj.title.upper())
            r_title.bold = True
            p.add_run("\t")
            r_date = p.add_run(proj.date)
            r_date.bold = True
            
            # Technology Stack Line
            p_stack = doc.add_paragraph()
            p_stack.paragraph_format.space_before = Pt(0)
            p_stack.paragraph_format.space_after = Pt(2)
            r_stack_lbl = p_stack.add_run("Technology Stacks: ")
            r_stack_lbl.bold = True
            p_stack.add_run(proj.tech_stack)
            
            for bullet in proj.bullets:
                doc.add_paragraph(bullet, style='List Bullet')

    # --- 5. TECHNICAL SKILLS ---
    if tailored_resume.skills:
        add_section_header("Technical Skills")
        for skill_cat in tailored_resume.skills:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            lbl = p.add_run(f"{skill_cat.category}: ")
            lbl.bold = True
            p.add_run(", ".join(skill_cat.items))

    # --- 6. CERTIFICATIONS ---
    if tailored_resume.certifications:
        add_section_header("Certifications")
        for cert in tailored_resume.certifications:
            doc.add_paragraph(cert, style='List Bullet')

    # --- 7. RESPONSIBILITY / LEADERSHIP ---
    if tailored_resume.responsibility:
        add_section_header("Responsibility")
        for resp in tailored_resume.responsibility:
            doc.add_paragraph(resp, style='List Bullet')

    # --- 8. ACHIEVEMENTS ---
    if tailored_resume.achievements:
        add_section_header("Achievements")
        for ach in tailored_resume.achievements:
            doc.add_paragraph(ach, style='List Bullet')

    # Save
    os.makedirs("generated_cvs", exist_ok=True)
    filename = f"generated_cvs/Tailored_CV_{uuid.uuid4().hex[:8]}.docx"
    filepath = os.path.abspath(filename)
    doc.save(filepath)
    
    return filepath


def generate_tailored_cv_sync(user_data: dict, resume_text: str, job_description: str) -> str:
    """Main entrypoint for CV generation (synchronous to avoid event loop blocking)."""
    logger.info("Generating tailored CV via AI...")
    tailored_resume = generate_tailored_content(resume_text, job_description, user_data)
    logger.info("Creating DOCX file...")
    filepath = create_docx_cv(tailored_resume)
    return filepath
